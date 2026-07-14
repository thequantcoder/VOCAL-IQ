#!/usr/bin/env python3
"""
Generate docs/PROJECT-FEATURES-EXPLAINED.docx from docs/PROJECT-FEATURES-EXPLAINED.md.

The Markdown file is the single source of truth. Whenever a new feature is added to the
project, append it to the .md (in the correct section), then run:

    python3 scripts/gen-features-docx.py

...to regenerate the beautifully-styled Word document. The two files always share the
same base name: PROJECT-FEATURES-EXPLAINED.

Only python-docx is required (already installed). No pandoc / network needed.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

# ---- Brand palette (VocalIQ: voice / waveform → indigo + sky accent) ----
PRIMARY = RGBColor(0x4F, 0x46, 0xE5)     # indigo-600
PRIMARY_HEX = "4F46E5"
ACCENT_HEX = "0EA5E9"                     # sky-500
HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)  # white on shaded header row
ALT_ROW_HEX = "EEF2FF"                    # indigo-50 (zebra stripe)
BORDER_HEX = "C7D2FE"                     # indigo-200
INK = RGBColor(0x1F, 0x29, 0x37)          # slate-800 body text
MUTED = RGBColor(0x6B, 0x72, 0x80)        # gray-500

REPO = Path(__file__).resolve().parent.parent
MD_PATH = REPO / "docs" / "PROJECT-FEATURES-EXPLAINED.md"
DOCX_PATH = REPO / "docs" / "PROJECT-FEATURES-EXPLAINED.docx"


# --------------------------------------------------------------------------- #
# Low-level docx helpers
# --------------------------------------------------------------------------- #
def _shade(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_borders(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), hex_color)
        borders.append(el)
    tcPr.append(borders)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def _add_inline(paragraph, text: str, *, bold=False, color=None, size=None, italic=False):
    """Add text to a paragraph, honouring **bold** and `code` markup."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        run = paragraph.add_run()
        run.bold = bold
        run.italic = italic
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0x9D, 0x17, 0x4D)
        else:
            run.text = part
        if color is not None:
            run.font.color.rgb = color
        if size is not None:
            run.font.size = Pt(size)


def _hr(doc):
    """Thin horizontal rule as a bottom-bordered empty paragraph."""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BORDER_HEX)
    borders.append(bottom)
    p_pr.append(borders)
    p.paragraph_format.space_after = Pt(4)


# --------------------------------------------------------------------------- #
# Markdown parsing (tailored to this document's structure)
# --------------------------------------------------------------------------- #
def parse_blocks(md: str):
    """Yield ('h1'|'h2'|'quote'|'para'|'bullet'|'table'|'hr', payload) blocks."""
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            yield ("hr", None)
            i += 1
            continue

        if stripped.startswith("# "):
            yield ("h1", stripped[2:].strip())
            i += 1
            continue
        if stripped.startswith("## "):
            yield ("h2", stripped[3:].strip())
            i += 1
            continue

        if stripped.startswith(">"):
            quote = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote.append(lines[i].strip().lstrip(">").strip())
                i += 1
            yield ("quote", quote)
            continue

        if stripped.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(lines[i].strip())
                i += 1
            # drop the |---|---| separator row
            cells = [
                [c.strip() for c in r.strip().strip("|").split("|")]
                for r in rows
                if not re.match(r"^\|[\s:|-]+\|?$", r)
            ]
            yield ("table", cells)
            continue

        if stripped.startswith("- "):
            bullets = []
            while i < len(lines) and lines[i].strip().startswith("- "):
                bullets.append(lines[i].strip()[2:].strip())
                i += 1
            yield ("bullet", bullets)
            continue

        # plain paragraph (single line in this doc)
        yield ("para", stripped)
        i += 1


# --------------------------------------------------------------------------- #
# Rendering
# --------------------------------------------------------------------------- #
def render_table(doc, cells):
    if not cells:
        return
    ncols = max(len(r) for r in cells)
    table = doc.add_table(rows=len(cells), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for r_idx, row in enumerate(cells):
        is_header = r_idx == 0
        for c_idx in range(ncols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.paragraphs[0].text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)

            if is_header:
                _shade(cell, PRIMARY_HEX)
                _add_inline(para, text, bold=True, color=HEADER_TEXT, size=10)
            else:
                if r_idx % 2 == 0:
                    _shade(cell, ALT_ROW_HEX)
                _add_inline(para, text, color=INK, size=10)
            _set_cell_borders(cell, BORDER_HEX)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def render_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    _add_inline(p, text, bold=True, color=PRIMARY, size=15)
    # accent underline
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), ACCENT_HEX)
    borders.append(bottom)
    p_pr.append(borders)


def render_quote(doc, quote_lines):
    for ln in quote_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(3)
        # subtle left bar
        p_pr = p._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), ACCENT_HEX)
        borders.append(left)
        p_pr.append(borders)
        _add_inline(p, ln, color=MUTED, size=10, italic=True)


def build(md: str) -> Document:
    doc = Document()

    # base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK

    first_h1_done = False
    for kind, payload in parse_blocks(md):
        if kind == "h1":
            if not first_h1_done:
                # cover title block
                title = doc.add_paragraph()
                title.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _add_inline(title, payload, bold=True, color=PRIMARY, size=26)
                sub = doc.add_paragraph()
                _add_inline(
                    sub,
                    "Har feature ka role + native / third-party / hybrid tag — ek jagah.",
                    color=MUTED, size=11, italic=True,
                )
                _hr(doc)
                first_h1_done = True
            else:
                render_h2(doc, payload)
        elif kind == "h2":
            render_h2(doc, payload)
        elif kind == "quote":
            render_quote(doc, payload)
        elif kind == "table":
            render_table(doc, payload)
        elif kind == "bullet":
            for b in payload:
                p = doc.add_paragraph(style="List Bullet")
                _add_inline(p, b, color=INK, size=10.5)
        elif kind == "hr":
            _hr(doc)
        elif kind == "para":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            _add_inline(p, payload, color=INK, size=10.5)

    # footer
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_inline(
        footer,
        "VocalIQ — Project Features Explained · auto-generated from PROJECT-FEATURES-EXPLAINED.md",
        color=MUTED, size=8,
    )
    return doc


def main() -> int:
    if not MD_PATH.exists():
        print(f"ERROR: source markdown not found: {MD_PATH}", file=sys.stderr)
        return 1
    md = MD_PATH.read_text(encoding="utf-8")
    doc = build(md)
    doc.save(DOCX_PATH)
    print(f"OK  wrote {DOCX_PATH.relative_to(REPO)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
